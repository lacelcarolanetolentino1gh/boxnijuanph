"""
Fixes two misalignments in all 4 BoxNiJuanPH submitted docs:
  1. "20 products" / "20 wellness products" → "27 products" / "27 wellness products"
  2. "10-keyword" → "15-keyword"
  3. "20 items" in catalog context → "27 products" (only where referring to product catalog size)
Preserves all formatting.
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import re
from docx import Document

FILES = [
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Technical-Proposal-Document-Revised.docx",
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Prototyping-Document (2).docx",
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Technical-Specification-Document-Revised (1).docx",
    r"C:\Users\I331629\Documents\BoxNiJuanPH-guide\BOXNIJUANPH-FINAL-PAPER.docx",
]

REPLACEMENTS = [
    # Product count — only where it refers to the catalog size
    (r"20 wellness products", "27 wellness products"),
    (r"20 products", "27 products"),
    # "from 20 items" / "from 20 wellness items" in context of catalog
    (r"from 20 items across", "from 27 products across"),
    (r"from 20 wellness items", "from 27 wellness products"),
    # "choosing from 20"
    (r"choosing from 20", "choosing from 27"),
    # "catalog of 20 items" / "catalog of 20 wellness products"
    (r"catalog of 20 items", "catalog of 27 products"),
    (r"catalog of 20 wellness products", "catalog of 27 wellness products"),
    # "20 items across 4 categories" in objectives
    (r"20 items across 4 categories", "27 products across 4 categories"),
    (r"20 items across four categories", "27 products across four categories"),
    # Standalone "20 items" only when followed by "across" (catalog ref)
    # ChatBot keyword count
    (r"10-keyword rule engine", "15-keyword rule engine"),
    (r"10-keyword", "15-keyword"),
    (r"10 keywords", "15 keywords"),
    (r"10 message types", "15 message types"),
]

def fix_run(run):
    changed = False
    for pattern, replacement in REPLACEMENTS:
        new_text = re.sub(pattern, replacement, run.text)
        if new_text != run.text:
            run.text = new_text
            changed = True
    return changed

def fix_para(para):
    changed = False
    for run in para.runs:
        if fix_run(run):
            changed = True
    return changed

def fix_doc(path):
    doc = Document(path)
    total = 0
    for para in doc.paragraphs:
        if fix_para(para): total += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if fix_para(para): total += 1
    doc.save(path)
    print(f"[OK] {path.split(chr(92))[-1]} — {total} paragraph(s) modified")

for f in FILES:
    try:
        fix_doc(f)
    except Exception as e:
        print(f"[ERROR] {f.split(chr(92))[-1]}: {e}")

print("\nAll done.")
