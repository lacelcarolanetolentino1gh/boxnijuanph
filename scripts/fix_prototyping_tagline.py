"""Patch the one remaining 'unlimited' in the Prototyping doc."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import re

path = r'C:\Users\I331629\Downloads\BoxNiJuanPH_Prototyping-Document (2).docx'
doc = Document(path)

count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if 'unlimited freedom' in run.text.lower():
            run.text = run.text.replace('unlimited freedom', 'complete flexibility')
            run.text = run.text.replace('Unlimited freedom', 'Complete flexibility')
            count += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if 'unlimited freedom' in run.text.lower():
                        run.text = run.text.replace('unlimited freedom', 'complete flexibility')
                        run.text = run.text.replace('Unlimited freedom', 'Complete flexibility')
                        count += 1

doc.save(path)
print(f'Done. {count} run(s) patched.')
