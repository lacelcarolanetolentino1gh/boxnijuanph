"""
Second-pass patch: catches remaining '20 wellness items' and '20 items' in catalog context.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import re
from docx import Document

FILES = [
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Technical-Proposal-Document-Revised.docx",
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Prototyping-Document (2).docx",
    r"C:\Users\I331629\Documents\BoxNiJuanPH-guide\BOXNIJUANPH-FINAL-PAPER.docx",
]

# More specific patterns missed in first pass
REPLACEMENTS = [
    (r"20 wellness items", "27 wellness products"),
    # "allowing users to build.*?choosing from.*?20" — catch "choosing from a curated catalog of 20"
    (r"curated catalog of 20 wellness products", "curated catalog of 27 wellness products"),
    (r"curated catalog of 20", "curated catalog of 27"),
    # "20 items" when describing catalog breadth in sentences
    (r"choosing from 20 items", "choosing from 27 products"),
    (r"browse.*?20 items", lambda m: m.group(0).replace("20 items", "27 products")),
    # "20 items across" — generic catch
    (r"\b20 items\b", "27 products"),
]

def fix_run(run):
    changed = False
    for item in REPLACEMENTS:
        pattern, replacement = item
        if callable(replacement):
            new_text = re.sub(pattern, replacement, run.text)
        else:
            new_text = re.sub(pattern, replacement, run.text)
        if new_text != run.text:
            run.text = new_text
            changed = True
    return changed

def fix_para(para):
    changed = False
    for run in para.runs:
        if fix_run(run):
            changed = True
    return changed

def fix_doc(path):
    doc = Document(path)
    total = 0
    for para in doc.paragraphs:
        if fix_para(para): total += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if fix_para(para): total += 1
    doc.save(path)
    print(f"[OK] {path.split(chr(92))[-1]} — {total} more paragraph(s) modified")

for f in FILES:
    try:
        fix_doc(f)
    except Exception as e:
        print(f"[ERROR] {f.split(chr(92))[-1]}: {e}")

print("\nDone.")
