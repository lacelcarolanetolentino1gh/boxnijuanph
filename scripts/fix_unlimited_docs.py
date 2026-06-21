"""
Replaces all instances of 'unlimited' / 'Unlimited' related to the Custom plan
with 'up to 12 items' phrasing across 4 BoxNiJuanPH documents.
Preserves all formatting (bold, color, font, size) via run-level replacement.
"""

import copy
import re
from docx import Document
from docx.oxml.ns import qn

FILES = [
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Technical-Proposal-Document-Revised.docx",
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Prototyping-Document (2).docx",
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Technical-Specification-Document-Revised (1).docx",
    r"C:\Users\I331629\Documents\BoxNiJuanPH-guide\BOXNIJUANPH-FINAL-PAPER.docx",
]

# (pattern, replacement) — case-insensitive regex on paragraph full text,
# but we do replacement at run level so formatting is preserved.
REPLACEMENTS = [
    # Table cell "Unlimited items" → "Up to 12 items"
    (r"[Uu]nlimited items", "Up to 12 items"),
    # "Unlimited" alone (e.g. in a standalone cell or badge)
    (r"^[Uu]nlimited$", "Up to 12 items"),
    # "Custom plan has no item limit" → "Custom plan allows up to 12 items per box"
    (r"[Cc]ustom plan has no item limit", "Custom plan allows up to 12 items per box"),
    # "no item limit" in isolation
    (r"no item limit", "up to 12 items per box"),
    # "Add as many items as you like" (builder banner copy in TSD)
    (r"Add as many items as you like", "Add up to 12 items to your box"),
    # "Custom (unlimited items)" parenthetical form
    (r"Custom \(unlimited items\)", "Custom (up to 12 items)"),
    # "Custom (unlimited items," with trailing comma
    (r"Custom \(unlimited items,", "Custom (up to 12 items,"),
    # "unlimited items," in running text
    (r"unlimited items,", "up to 12 items,"),
    # "unlimited items" in running text (catch-all, after specific patterns)
    (r"unlimited items", "up to 12 items"),
    # "Unlimited" standalone at start of cell (e.g. table "Items per Box" column)
    (r"^Unlimited$", "Up to 12 items"),
]


def replace_in_run(run, pattern, replacement):
    """Replace regex pattern in a single run's text, preserving the run's XML formatting."""
    new_text = re.sub(pattern, replacement, run.text)
    if new_text != run.text:
        run.text = new_text
        return True
    return False


def process_paragraph(para):
    changed = False
    for run in para.runs:
        for pattern, replacement in REPLACEMENTS:
            if replace_in_run(run, pattern, replacement):
                changed = True
    return changed


def process_document(path):
    doc = Document(path)
    total_changes = 0

    # Body paragraphs
    for para in doc.paragraphs:
        if process_paragraph(para):
            total_changes += 1

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph(para):
                        total_changes += 1

    # Headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    if process_paragraph(para):
                        total_changes += 1

    doc.save(path)
    print(f"[OK] {path}")
    print(f"     {total_changes} paragraph(s) modified\n")


for f in FILES:
    try:
        process_document(f)
    except Exception as e:
        print(f"[ERROR] {f}\n  {e}\n")

print("All done.")
