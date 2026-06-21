"""
Fixes the remaining 'Product Catalog — 20 Items Across 4 Categories' heading
in the two docs that still have it.
Iterates ALL paragraphs (including headings) and replaces at the run level.
Also does a full-text fallback: if run-level replacement misses it (split runs),
rebuilds the paragraph text.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

FILES = [
    r"C:\Users\I331629\Downloads\BoxNiJuanPH_Technical-Proposal-Document-Revised.docx",
    r"C:\Users\I331629\Documents\BoxNiJuanPH-guide\BOXNIJUANPH-FINAL-PAPER.docx",
]

OLD = "20 Items Across 4 Categories"
NEW = "27 Products Across 4 Categories"

def fix_para(para):
    # First try run-level replacement
    for run in para.runs:
        if OLD in run.text:
            run.text = run.text.replace(OLD, NEW)
            return True
    # Fallback: check full paragraph text (handles split runs)
    full = "".join(r.text for r in para.runs)
    if OLD in full:
        # Replace in first run that contains any part, rebuild
        # Simplest safe approach: clear all runs and put text in first run
        new_full = full.replace(OLD, NEW)
        if para.runs:
            para.runs[0].text = new_full
            for run in para.runs[1:]:
                run.text = ""
        return True
    return False

def fix_doc(path):
    doc = Document(path)
    total = 0
    for para in doc.paragraphs:
        if fix_para(para):
            total += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if fix_para(para):
                        total += 1
    doc.save(path)
    print(f"[OK] {path.split(chr(92))[-1]} — {total} paragraph(s) fixed")

for f in FILES:
    try:
        fix_doc(f)
    except Exception as e:
        print(f"[ERROR] {f.split(chr(92))[-1]}: {e}")

print("\nDone.")
