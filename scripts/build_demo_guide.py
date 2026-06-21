"""
Generates BoxNiJuanPH_DemoGuide_June21.docx
A step-by-step click-by-click demo cheat sheet for Lacel to follow during the live prototype demo.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

GREEN  = RGBColor(0x5F, 0x8F, 0x72)
DARK   = RGBColor(0x2D, 0x2D, 0x2D)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
AMBER  = RGBColor(0x92, 0x40, 0x0E)
RED    = RGBColor(0xDC, 0x26, 0x26)
PURPLE = RGBColor(0x6D, 0x28, 0xD9)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

for section in doc.sections:
    section.top_margin    = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── helpers ───────────────────────────────────────────────────────

def dark_banner(text, fill="2D2D2D", size=11.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    r = p.add_run("  " + text + "  ")
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = WHITE

def green_banner(text, fill="EAF2ED", text_color=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    r = p.add_run("  " + text + "  ")
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = text_color or GREEN

def step(num, action, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(14)
    r1 = p.add_run(f"{num}.  ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = GREEN
    r2 = p.add_run(action)
    r2.font.size = Pt(11); r2.font.color.rgb = DARK
    if note:
        r3 = p.add_run(f"  ← {note}")
        r3.italic = True; r3.font.size = Pt(10); r3.font.color.rgb = GRAY

def say(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Pt(28)
    r1 = p.add_run("SAY:  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = PURPLE
    r2 = p.add_run(f'"{text}"')
    r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = PURPLE

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(14)
    r1 = p.add_run("⚠  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = AMBER
    r2 = p.add_run(text)
    r2.font.size = Pt(10); r2.font.color.rgb = AMBER; r2.italic = True

def checkpoint(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(14)
    r1 = p.add_run("✓  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = GREEN
    r2 = p.add_run(text)
    r2.font.size = Pt(10); r2.font.color.rgb = GREEN; r2.bold = True

def divider(color="D1FAE5", size=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), str(size))
    bot.set(qn('w:space'), '1');   bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def blank(space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)

# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("BoxNiJuanPH")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("DEMO GUIDE — Lacel's Click-by-Click Cheat Sheet")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("June 21, 2026  ·  COMP 047  ·  boxnijuanph.vercel.app")
r.font.size = Pt(10); r.italic = True; r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
r = p.add_run("Keep this open beside your browser during the demo. Follow each step in order.")
r.font.size = Pt(10); r.font.color.rgb = AMBER; r.bold = True

divider("5F8F72", 6)

# Before you start box
green_banner("BEFORE YOU START — do these BEFORE the presentation begins", fill="FFF7ED", text_color=AMBER)
step("A", "Open boxnijuanph.vercel.app in your browser (regular window, not incognito)")
step("B", "Sign in as a test user — so My Box dashboard is ready to show")
step("C", "Pre-fill the box with 5 Standard items and complete a checkout — so /my-box shows an active subscription")
step("D", "Keep this guide open on your phone or a second window/tab")
warn("Do NOT clear localStorage or refresh mid-demo — it will wipe your test data.")
warn("If you make a mistake, keep talking and scroll past it calmly. The professor won't know.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 1 — HOMEPAGE
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 1 of 8 — HOMEPAGE  (/)")
green_banner("Slide 11 is showing. You are now live on the site.", fill="EAF2ED")

step(1, "Make sure you're on the homepage (/).")
step(2, "Point to the HERO section — headline, trust badges, two buttons.")
say("Ito po ang Homepage namin. Dito makikita ang hero section at ang trust badges sa ibaba.")

step(3, "Scroll down slowly — show How It Works (4 steps).")
say("Ang How It Works section — apat na steps mula sa pagpili ng plan hanggang matanggap ang box.")

step(4, "Scroll to the WELLNESS QUIZ section (dark background).")
say("May Wellness Quiz din kami — tatlong tanong lang, irerecommend nito ang tamang plan.")
step(5, "Click any answer on Q1 (e.g. 'Moderate').")
step(6, "Click any answer on Q2 (e.g. 'All-Around').")
step(7, "Click 'Load me up' on Q3.")
checkpoint("Result screen shows: Premium Plan, ₱899/mo")
say("Nakita ninyo — base sa sagot, irerecommend ang Premium Plan.")
step(8, "Click 'Retake Quiz' — quiz resets. Do NOT click 'Get This Plan' yet.")

blank(4)
step(9, "Scroll down — briefly show Featured Products, Plans Preview, Testimonials sections.")
say("Dito rin makikita ang featured products, plans preview, at testimonials mula sa aming subscribers.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 2 — PLANS PAGE
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 2 of 8 — PLANS PAGE  (/plans)")
green_banner("Click 'Build Your Box' in the navbar OR scroll to CTA and click 'Get Started Today →'")

step(1, "Navigate to /plans.")
say("Pupunta tayo sa Plans page. Makikita ninyo ang apat na subscription tiers.")

step(2, "Point to each plan card — Basic, Standard, Premium, Custom.")
say("Basic sa ₱399, Standard sa ₱599, Premium sa ₱899, at Custom sa ₱1,299.")

step(3, "Point to the SAVINGS BADGE on any card — the strikethrough price + green pill.")
say("Ang savings badge — nagpapakita kung magkano ang nati-tipid kumpara sa one-time purchase.")

step(4, "Scroll down slightly — show the LOYALTY PERKS STRIP (light green box).")
say("At dito — ang subscriber perks: Birthday Voucher, Loyalty Points, Subscriber-Only Deals, at Pause Anytime.")

step(5, "Click 'Choose Standard' — navigate to the builder with Standard plan.")
warn("Make sure to click Standard, not Basic or Premium.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 3 — BOX BUILDER
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 3 of 8 — BOX BUILDER  (/builder)")
green_banner("You should now be on /builder with 'Standard Plan — Choose 5 items'")

step(1, "Point to the heading and progress bar.")
say("Ito ang Box Builder. Standard plan — kailangan ng 5 items. Makikita ang progress bar.")

step(2, "Type 'foam' in the search bar — show filtering.")
say("May search at filters — category, brand, at Local brands only.")
step(3, "Clear the search (click X).")

step(4, "Click the ⓘ button on any product card (small circle, bottom-right of the image).")
say("Ang ⓘ button — makikita ang full product details nang hindi pa nase-select ang item.")
checkpoint("Info modal opens. Product is NOT selected (no green border).")
step(5, "Close the modal — click 'Got it — back to building'.")

step(6, "Click a product card to select it — Variant Modal appears.")
say("Kapag nag-click ng card, lalabas ang Variant Modal para pumili ng specific variant.")
step(7, "Select a variant. Click 'Add to Box →'.")
checkpoint("Card now has green border + checkmark. Progress bar updates to 1/5.")

step(8, "Add 4 more products the same way until box is full (5/5).")
checkpoint("Progress bar shows '✓ Box complete!' and remaining cards go gray.")
say("Puno na ang box — 5 items na.")

step(9, "Point to the right sidebar showing all 5 selected items.")
say("Dito sa sidebar — makikita ang lahat ng napiling items.")

step(10, "Click 'Review My Box →'.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 4 — BOX SUMMARY
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 4 of 8 — BOX SUMMARY  (/summary)")
green_banner("You should now be on /summary showing all 5 selected items")

step(1, "Point to the plan banner at the top — plan name, price.")
say("Ang Box Summary — review ng lahat ng items bago mag-checkout.")

step(2, "Point to the item list — show Local and Eco badges on applicable items.")
say("Makikita ang Filipino Brand at Eco-Friendly badges sa mga applicable na products.")

step(3, "Point to the CSR IMPACT NOTE card.")
say("At dito — real-time count ng local brands at eco-friendly items sa box nila. Aligned sa UN SDG 12.")

step(4, "Click 'Proceed to Checkout →'.")
warn("Do NOT click '← Edit Box' — you'll go back to the builder.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 5 — LOGIN / SIGN IN
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 5 of 8 — SIGN IN  (/login or /checkout)")
green_banner("If you pre-logged in before the demo, you'll skip straight to checkout — that's fine.")

step(1, "If a login page appears — sign in with your test account.")
step(2, "If checkout shows the green 'Signed in as [Name]' chip — you're already logged in, skip to Step 6.")
say("Naka-sign in na kami — makikita ang green chip dito sa taas.")
warn("If you see the amber GUEST banner — just proceed as guest. It still works.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 6 — CHECKOUT
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 6 of 8 — CHECKOUT  (/checkout)")
green_banner("You should now be on /checkout with the order summary visible on the right")

step(1, "Point to the delivery form on the left.")
say("Ang checkout form — may built-in validation para sa phone number at address.")

step(2, "Point to phone field — briefly type '12345' to show validation.")
checkpoint("Red error appears: 'Enter a valid PH number'.")
step(3, "Clear and type '09171234567'.")
checkpoint("Error disappears. Valid format accepted.")
say("Philippine phone number format lang ang tinatanggap.")

step(4, "Fill in the rest of the form — Full Name, Email, Address, City, ZIP.")
warn("Use fake data — e.g. Name: Juan dela Cruz, Email: juan@test.com, Address: 123 Rizal St, City: Manila, ZIP: 1000")

step(5, "Select a payment method — click GCash.")
step(6, "Point to the ORDER SUMMARY sidebar on the right.")
say("Dito sa kanan — makikita ang order summary: lahat ng items at ang plan price.")

step(7, "Click 'Place Order →'.")
checkpoint("Confirmation dialog appears showing delivery details + plan + price.")
say("May confirmation dialog bago ma-finalize ang order.")
step(8, "Click 'Confirm Order'.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 7 — ORDER CONFIRMATION
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 7 of 8 — ORDER CONFIRMATION  (/confirmation)")
green_banner("You should now be on /confirmation with a green checkmark")

step(1, "Point to the green checkmark and 'Order Confirmed!' heading.")
say("Order Confirmed! May order number sa BNJ-XXXXXX format.")

step(2, "Point to the order number (monospace font).")
step(3, "Point to the item list — show the Local and Eco badges.")
step(4, "Point to the CSR THANK-YOU section.")
say("Salamat sa iyong suporta sa mga lokal na wellness brand ng Pilipinas.")

step(5, "Scroll down — show the AMBER REFUND NOTICE CARD.")
say("At dito — nagpapaalam sa user na may 7-day window sila para mag-request ng refund o replacement.")

step(6, "Click 'View My Subscription →'.")

divider()

# ══════════════════════════════════════════════════════════════════
# STEP 8 — MY BOX
# ══════════════════════════════════════════════════════════════════
dark_banner("STEP 8 of 8 — MY BOX DASHBOARD  (/my-box)")
green_banner("You should now be on /my-box showing the active subscription")

step(1, "Point to the plan card — green top bar, ✓ Active status, plan name and price.")
say("Ang My Box Dashboard — dito mina-manage ng user ang kanilang subscription.")

step(2, "Point to the box contents list — all 5 items from the order.")
say("Makikita ang lahat ng items sa box — pati ang Local at Eco badges.")

step(3, "Scroll down to the MANAGE section.")
step(4, "Point to '↩️ Report an Issue' button.")
say("Kung may problema sa order — isang click lang para pumunta sa contact form na pre-filled na ang topic.")

step(5, "Point to '⏸ Pause Subscription' button.")
say("Pwede ring i-pause ang subscription — walang extra charge, walang lock-in.")

step(6, "Click the '👤 Profile' tab.")
step(7, "Point to the Personal Information section and the Default Payment Method section.")
say("At dito ang Profile tab — editable ang personal info at default payment method.")

blank(4)
checkpoint("DEMO COMPLETE — lahat ng 8 steps tapos na.")
say("Iyon po ang buong platform — live, fully functional, zero cost. Maraming salamat po!")

divider("5F8F72", 6)

# ══════════════════════════════════════════════════════════════════
# QUICK RECOVERY GUIDE
# ══════════════════════════════════════════════════════════════════
dark_banner("IF SOMETHING GOES WRONG — Quick Recovery", fill="92400E")

green_banner("Page shows blank / error", fill="FFF7ED", text_color=AMBER)
step(1, "Press Ctrl+Shift+R (hard refresh). Keep talking while it loads.")
step(2, "If still broken — navigate manually: type the URL in the address bar.")

green_banner("You clicked the wrong button and went to the wrong page", fill="FFF7ED", text_color=AMBER)
step(1, "Don't use browser Back — use the in-app navigation links instead.")
step(2, "E.g. if you're on /builder by mistake — click '← Back to Plans' inside the page.")

green_banner("Items disappeared from the box / My Box shows empty", fill="FFF7ED", text_color=AMBER)
step(1, "This means localStorage was cleared. Skip to My Box and say:")
say("Para sa demo — pupunta na tayo directly sa My Box para makita ang subscription dashboard.")

green_banner("ChatBot won't open or respond", fill="FFF7ED", text_color=AMBER)
step(1, "Click the green bubble at the bottom-right.")
step(2, "If no response in 2 seconds — type 'refund' and click Send.")
step(3, "If still broken — skip it. Say: 'Ang BoxBot ay available sa lahat ng pages for live support.'")

divider()

# ══════════════════════════════════════════════════════════════════
# WHAT TO HIGHLIGHT PER SLIDE
# ══════════════════════════════════════════════════════════════════
dark_banner("SLIDE-TO-DEMO REFERENCE — What to point to on each slide")

rows = [
    ("Slide 11", "Homepage & Plans", "Hero → Wellness Quiz → /plans savings badges → loyalty perks"),
    ("Slide 12", "Builder & Summary", "ⓘ button → Variant Modal → progress bar → CSR count on /summary"),
    ("Slide 13", "Checkout & Confirmation", "Phone validation → confirmation dialog → amber refund card → /my-box"),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
for cell, txt in zip(table.rows[0].cells, ["Slide", "Section", "Key things to show"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = WHITE
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '2D2D2D')
    tcPr.append(shd)

for slide_num, section, what in rows:
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, [slide_num, section, what])):
        cell.text = val
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK
        if i == 0:
            run.bold = True; run.font.color.rgb = GREEN

blank(10)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BoxNiJuanPH  ·  COMP 047  ·  PUP Open University System  ·  June 2026  ·  Lacel's Demo Guide")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY

out = r"C:\Users\I331629\Documents\BoxNiJuanPH-guide\BoxNiJuanPH_DemoGuide_June21.docx"
doc.save(out)
print("Done:", out)
