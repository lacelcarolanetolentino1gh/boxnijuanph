"""
Converts BoxNiJuanPH_PresentationScript_June21.txt to a formatted Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

GREEN  = RGBColor(0x5F, 0x8F, 0x72)
DARK   = RGBColor(0x2D, 0x2D, 0x2D)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
AMBER  = RGBColor(0x92, 0x40, 0x0E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLUE   = RGBColor(0x1D, 0x4E, 0xD8)

PRESENTER_COLORS = {
    "RIZZA":  RGBColor(0x5F, 0x8F, 0x72),   # green
    "CARLA":  RGBColor(0x1D, 0x4E, 0xD8),   # blue
    "JULIO":  RGBColor(0x92, 0x40, 0x0E),   # amber
    "LACEL":  RGBColor(0x6D, 0x28, 0xD9),   # purple
}

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── helpers ───────────────────────────────────────────────────────

def shaded_para(text, fill_hex, text_color, bold=True, size=11):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("  " + text + "  ")
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = text_color
    return p

def divider(color_hex="E5E7EB", size=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def body_line(text, color=None, bold=False, italic=False, indent=0, size=10.5, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(text)
    r.bold    = bold
    r.italic  = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or DARK
    return p

def blank(space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space)

# ── COVER ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(28)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("BoxNiJuanPH")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Oral Presentation Script — Taglish (More English)")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("COMP 047  ·  BSITOUMN 2-1  ·  PUP Open University System  ·  June 21, 2026")
r.font.size = Pt(10); r.italic = True; r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Presenters: Lacel · Carla · Julio · Rizza  (Ivy — absent, power interruption)")
r.font.size = Pt(10); r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("Estimated Total Time: 13–15 minutes")
r.font.size = Pt(10); r.bold = True; r.font.color.rgb = GREEN

divider("5F8F72", 6)

# Legend note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("Color guide:  ")
r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = DARK

for name, col in PRESENTER_COLORS.items():
    r = p.add_run(f"■ {name}   ")
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = col

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(10)
r = p2.add_run("(Stage directions are in italic gray  ·  [SLIDE X] = advance the slide)")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRAY

divider()

# ── SLIDES ────────────────────────────────────────────────────────

SLIDES = [
    {
        "num": 1, "title": "TITLE", "presenter": "RIZZA", "time": "~30 seconds",
        "direction": "[Advance to Slide 1]",
        "lines": [
            ("Good morning (or good afternoon), Professor Carabaca!", False, False),
            ("We are BSITOUMN 2-1, presenting our final project for COMP 047.", False, False),
            ("It is our honor to introduce:", False, False),
            ('"BoxNiJuanPH — An Interactive Wellness Subscription Box Platform.', False, False),
            ('The Wellness Box for Every Juan."', False, False),
            ("Our platform is live and fully accessible right now\nat boxnijuanph.vercel.app.", False, False),
            ("Let's get started.", False, False),
        ],
    },
    {
        "num": 2, "title": "PRESENTATION OUTLINE", "presenter": "RIZZA", "time": "~30 seconds",
        "direction": "[Advance to Slide 2]",
        "lines": [
            ("Our presentation today is structured into five parts:", False, False),
            ("First — Introduction and Background of the problem.", False, False),
            ("Second — our six SMART Objectives.", False, False),
            ("Third — a comparison with existing related systems.", False, False),
            ("Fourth — the Technical Specifications of our platform.", False, False),
            ("And finally — a live Prototype Demonstration, which Lacel will walk us through.", False, False),
            ("Before we begin — one of our members, Ivy, is unable to join today\ndue to a power interruption in her area.\nHer assigned sections have been redistributed among the group.", False, False),
        ],
    },
    {
        "num": 3, "title": "INTRODUCTION & BACKGROUND", "presenter": "CARLA", "time": "~2 minutes",
        "direction": "[Advance to Slide 3]",
        "lines": [
            ("Let's start with the problem.", False, False),
            ("When you subscribe to a wellness subscription box —\ndo you actually get to choose what's inside?", False, False),
            ("In most cases, no. Everything is pre-curated.\nThe service provider decides what you receive.\nYou are never asked about your preferences.", False, False),
            ("The result? Users end up with products they didn't want.\nThat leads to wasted spending and high cancellation rates.", False, False),
            ("On top of that — the available platforms like Cratejoy and FabFitFun\nare US-based. There's no Philippine Peso pricing.\nNo GCash, no Maya, and no local Filipino brands.", False, False),
            ("So Filipino consumers — particularly the health-conscious\n18 to 35 age group — have no platform that truly fits their needs.", False, False),
            ("The market opportunity, however, is significant.\nAccording to Grand View Research 2023, the global subscription box market\nwas valued at 32 billion dollars in 2022\nand is projected to reach 105 billion by 2032 —\ngrowing at an 18.3% compound annual growth rate.", False, False),
            ("And according to Epsilon 2022, 80% of consumers are more likely to purchase\nfrom brands that offer personalized experiences.", False, False),
            ("So we asked: what if we built something different?", False, False),
            ("BoxNiJuanPH flips the model entirely —\nYOU choose every single item in your monthly wellness box.", True, False),
        ],
    },
    {
        "num": 4, "title": "SMART OBJECTIVES", "presenter": "CARLA", "time": "~1.5 minutes",
        "direction": "[Advance to Slide 4]",
        "lines": [
            ("Our project is anchored on six SMART Objectives:", False, False),
            ("O1 — Full personalization.\nUsers can build and order a customized box\nfrom 27 wellness products across four categories — delivered by June 14, 2026.", False, False),
            ("O2 — Four subscription tiers with clear pricing:\nBasic at \u20b1399, Standard at \u20b1599,\nPremium at \u20b1899, and Custom at \u20b11,299 —\nall competitively priced for the Philippine market.", False, False),
            ("O3 — An intuitive user experience designed for Filipino consumers aged 18 to 35,\nregardless of their technical background.", False, False),
            ("O4 — Integrated CSR — with Filipino Brand tags, Eco-Friendly badges,\nand a real-time CSR impact summary after every order,\naligned with UN Sustainable Development Goal 12:\nResponsible Consumption and Production.", False, False),
            ("O5 — Credible market positioning supported by four or more cited research references.", False, False),
            ("O6 — Complete and on-time submission — full working platform,\nall required documentation, and this presentation —\ndelivered by June 14, 2026.", False, False),
            ("We are happy to report that all six objectives have been met.", True, False),
        ],
    },
    {
        "num": 5, "title": "RELATED SYSTEMS", "presenter": "JULIO", "time": "~1.5 minutes",
        "direction": "[Advance to Slide 5]",
        "lines": [
            ("We conducted a comparative analysis of existing platforms\nto identify the gaps that BoxNiJuanPH addresses.", False, False),
            ("Cratejoy and FabFitFun are US-based.\nNo PHP pricing, no local brand support,\nand very limited customization options.", False, False),
            ("Good Box PH operates locally, but their boxes are fully pre-curated.\nThere is no item-level customization,\nand no subscription management dashboard.", False, False),
            ("Boozy.ph and BeautyMNL are niche platforms —\nwellness subscription boxes are not their primary focus.", False, False),
            ("BoxNiJuanPH fills that gap.\nIt is the only platform in the Philippine market that offers:\nfull item-level customization across all plan tiers,\ncurated Filipino brand integration, PHP pricing,\nsubscription management tools, and built-in CSR features —", False, False),
            ("all in one platform, built specifically for Filipinos.", True, False),
        ],
    },
    {
        "num": 6, "title": "TECH SPECS: ARCHITECTURE", "presenter": "JULIO", "time": "~1 minute",
        "direction": "[Advance to Slide 6]",
        "lines": [
            ("Our entire tech stack was built at zero cost.", False, False),
            ("We used Next.js App Router with TypeScript\nand Tailwind CSS version 4 for the frontend.\nThe application is deployed on Vercel's free Hobby Plan.", False, False),
            ("All user data is stored locally on the user's device.\nThere is no external server and no third-party database.\nThis architecture ensures full compliance with RA 10173 —\nthe Data Privacy Act of 2012 — right out of the box.", False, False),
            ("The platform is live right now at boxnijuanph.vercel.app.\nTotal infrastructure cost: zero pesos.", True, False),
        ],
    },
    {
        "num": 7, "title": "TECH SPECS: PAGES & FEATURES", "presenter": "JULIO", "time": "~1 minute",
        "direction": "[Advance to Slide 7]",
        "lines": [
            ("BoxNiJuanPH covers the complete end-to-end user journey\nacross 11 fully responsive pages.", False, False),
            ("Starting from the Homepage — where users can explore featured products,\nbrowse a plans preview, and take our three-question Wellness Quiz\nthat recommends the right plan based on their lifestyle.", False, False),
            ("The /plans page presents our four subscription tiers.", False, False),
            ("The /builder page is the interactive box builder —\nfeaturing 27 products, category and brand filters,\nvariant selection, and a real-time progress bar.", False, False),
            ("The /my-box page is the subscription dashboard —\nwhere users manage their active subscription, update their box,\nor pause their subscription at any time.", False, False),
            ("And our Custom plan offers up to 12 items per box —\nthe most flexible tier we offer.", False, False),
        ],
    },
    {
        "num": 8, "title": "USE CASES", "presenter": "JULIO", "time": "~45 seconds",
        "direction": "[Advance to Slide 8]",
        "lines": [
            ("We documented four use cases covering the key user scenarios.", False, False),
            ("UC1 — A registered user on the Standard plan.\nThey select 5 items, complete checkout with a pre-filled form,\nreceive an order number, and land on an active subscription in My Box.", False, False),
            ("UC2 — Guest checkout.\nLimitations are clearly flagged with an amber notice,\nthey complete the manual form, and the order is placed.", False, False),
            ("UC3 — A returning user.\nThey update their profile and set a default GCash payment method —\nwhich auto-fills on their next checkout.", False, False),
            ("UC4 — Cancelling a subscription.\nJust two confirmation steps.\nNo lock-in period. No phone call required.", False, False),
        ],
    },
    {
        "num": 9, "title": "CSR & SUSTAINABILITY", "presenter": "CARLA", "time": "~1 minute",
        "direction": "[Advance to Slide 9]",
        "lines": [
            ("We want to emphasize — our CSR integration is not decorative.\nIt is embedded into the core purchase experience.", False, False),
            ("First — every locally sourced product carries a Filipino Brand badge.\nUsers see this badge while browsing and on every order summary.", False, False),
            ("Second — because users personally choose every item in their box,\nthere are no unwanted products and no unnecessary waste.\nThis is a direct application of UN SDG 12 —\nResponsible Consumption and Production.", False, False),
            ("Third — on the Box Summary, Order Confirmation, and My Box pages,\nusers see a real-time count of how many local and eco-friendly items\nare in their box. Not just a label — an actual running total.", False, False),
            ("This is our commitment to responsible e-commerce\nthat goes beyond aesthetics.", True, False),
        ],
    },
    {
        "num": 10, "title": "PROTOTYPE DEMO — INTRO", "presenter": "LACEL", "time": "~30 seconds",
        "direction": "[Advance to Slide 10 — Lacel takes screen]",
        "lines": [
            ("Good afternoon again, Professor!", False, False),
            ("I'm Lacel — Project Lead and UI/UX Designer of BoxNiJuanPH.", False, False),
            ("I'll now walk you through the full platform — live —\nat boxnijuanph.vercel.app.", False, False),
            ("Our demo covers eight steps:\nHomepage, Plan Selection, Box Builder, Box Summary,\nSign In, Checkout, Order Confirmation,\nand the My Box Dashboard.", False, False),
            ("Let's begin.", True, False),
        ],
    },
    {
        "num": 11, "title": "DEMO: HOMEPAGE & PLANS", "presenter": "LACEL", "time": "~1.5 minutes",
        "direction": "[Advance to Slide 11 — show live site]",
        "lines": [
            ("This is our Homepage.", False, False),
            ("You can see the hero section here —\nwith the headline \"Your Personal Wellness Box Every Month\",\ntrust badges below it, and the How It Works section\nthat walks new users through the four-step process.", False, False),
            ("We also have a Wellness Quiz on this page —\njust three questions, and it recommends the right subscription plan\nbased on the user's lifestyle and preferences.", False, False),
            ("(Demo the quiz briefly — click through all 3 questions)", False, True),
            ("Now let's head to the Plans page.", False, False),
            ("Here you can see all four subscription tiers:\nBasic, Standard, Premium, and Custom.", False, False),
            ("Notice the savings badges on each card —\nthey show exactly how much the user saves\ncompared to buying the same products individually.", False, False),
            ("And down here — our loyalty perks strip:\nBirthday Box Voucher, Loyalty Points,\nSubscriber-Only Deals, and the option to Pause Anytime.", False, False),
        ],
    },
    {
        "num": 12, "title": "DEMO: BOX BUILDER & SUMMARY", "presenter": "LACEL", "time": "~2 minutes",
        "direction": "[Advance to Slide 12 — show live builder]",
        "lines": [
            ("Let's click on Standard Plan and go into the Box Builder.", False, False),
            ("Here you can see our 27 products,\nalong with category and brand filters.\nUsers can also filter specifically for local brands only.", False, False),
            ("(Select a product — show variant modal)", False, True),
            ("When you click on a product card,\nthe Variant Modal appears — the user picks their preferred variant\nand adds it to their box.", False, False),
            ("Notice the green border and checkmark that appears on the card\nonce an item is added — and the progress bar updates in real time.", False, False),
            ("(Show \u24d8 button)", False, True),
            ("There's also this \u24d8 button on every card.\nClick it, and you get the full product details —\nwithout selecting or committing to the item.\nThis is especially useful on mobile.", False, False),
            ("(Fill the box and go to /summary)", False, True),
            ("Once the box is full, we move to the Box Summary page\nto review all selected items before checkout.", False, False),
            ("And here is the CSR impact count —\nshowing how many local brands and eco-friendly items the user chose.", False, False),
        ],
    },
    {
        "num": 13, "title": "DEMO: CHECKOUT & CONFIRMATION", "presenter": "LACEL", "time": "~1.5 minutes",
        "direction": "[Advance to Slide 13 — proceed to checkout]",
        "lines": [
            ("Proceeding to Checkout.", False, False),
            ("The delivery form has built-in validation —\nPhilippine phone number format is enforced,\naddress fields have minimum length requirements,\nand all required fields must be completed before submission.", False, False),
            ("(Show order confirmation dialog, then confirm)", False, True),
            ("When the form is complete and the user clicks \"Place Order\",\na confirmation dialog appears to review the full order\nbefore it's finalized.", False, False),
            ("(Show /confirmation page)", False, True),
            ("And this is the Order Confirmation page.", False, False),
            ("The user sees their order number in BNJ-XXXXXX format,\nthe full item list with local and eco badges,\nthe CSR thank-you message,\nand this amber refund notice card —\ninforming the user that they have a 7-day window\nto request a refund or replacement.", False, False),
            ("(Navigate to /my-box)", False, True),
            ("And finally — the My Box Dashboard.", False, False),
            ("This is where users manage their active subscription:\nthey can view their box items, switch their plan,\npause their subscription, or report an issue —\nwhich takes them directly to a pre-filled contact form.", False, False),
        ],
    },
    {
        "num": 14, "title": "CLOSING / Q&A", "presenter": "RIZZA", "time": "~30 seconds",
        "direction": "[Advance to Slide 14]",
        "lines": [
            ("That is BoxNiJuanPH —\na fully functional, live, zero-cost, Filipino-first\nwellness subscription box platform.", False, False),
            ("All six SMART objectives have been accomplished.\nThe platform is deployed and accessible right now\nat boxnijuanph.vercel.app.", False, False),
            ("We would like to thank you, Professor Carabaca,\nfor the guidance and feedback throughout this course.", False, False),
            ("We are now open for questions.", False, False),
            ("And if you'd like to try the platform yourself —\nit takes about five minutes to complete a full order end to end.", False, False),
            ("Maraming salamat po!", True, False),
        ],
    },
]

FILL_MAP = {
    "RIZZA": "EAF2ED",   # light green
    "CARLA": "EFF6FF",   # light blue
    "JULIO": "FFFBEB",   # light amber
    "LACEL": "F5F3FF",   # light purple
}
TEXT_MAP = {
    "RIZZA": GREEN,
    "CARLA": BLUE,
    "JULIO": AMBER,
    "LACEL": RGBColor(0x6D, 0x28, 0xD9),
}

for slide in SLIDES:
    pres   = slide["presenter"]
    fill   = FILL_MAP[pres]
    tcolor = TEXT_MAP[pres]

    # Slide header band
    shaded_para(
        f"SLIDE {slide['num']} — {slide['title']}   ({pres}, {slide['time']})",
        fill_hex="2D2D2D",
        text_color=WHITE,
        size=10.5,
    )

    # Stage direction
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(10)
    r = p.add_run(slide["direction"])
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GRAY

    # Presenter label
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    r = p.add_run(f"  {pres}:  ")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = tcolor

    # Lines
    for (text, bold, is_direction) in slide["lines"]:
        for subline in text.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Pt(18)
            r = p.add_run(subline)
            r.font.size = Pt(10.5)
            if is_direction:
                r.italic = True
                r.font.color.rgb = GRAY
            elif bold:
                r.bold = True
                r.font.color.rgb = tcolor
            else:
                r.font.color.rgb = DARK

    blank(8)
    divider("5F8F72" if slide["num"] == 14 else "D1FAE5", 4 if slide["num"] == 14 else 2)

# ── TIMING TABLE ──────────────────────────────────────────────────
blank(10)
shaded_para("TIMING GUIDE", fill_hex="2D2D2D", text_color=WHITE, size=11)

from docx.shared import Inches as In
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as nqn

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["Slide", "Presenter", "Section", "Est. Time"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = WHITE
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(nqn('w:val'), 'clear'); shd.set(nqn('w:color'), 'auto'); shd.set(nqn('w:fill'), '2D2D2D')
    tcPr.append(shd)

TIMING = [
    ("1",  "Rizza", "Title",                  "~0:30"),
    ("2",  "Rizza", "Outline + Ivy note",      "~0:30"),
    ("3",  "Carla", "Introduction & Background","~2:00"),
    ("4",  "Carla", "SMART Objectives",        "~1:30"),
    ("5",  "Julio", "Related Systems",         "~1:30"),
    ("6",  "Julio", "Architecture",            "~1:00"),
    ("7",  "Julio", "Pages & Features",         "~1:00"),
    ("8",  "Julio", "Use Cases",               "~0:45"),
    ("9",  "Carla", "CSR & Sustainability",    "~1:00"),
    ("10", "Lacel", "Demo Intro",              "~0:30"),
    ("11", "Lacel", "Homepage & Plans demo",   "~1:30"),
    ("12", "Lacel", "Builder & Summary demo",  "~2:00"),
    ("13", "Lacel", "Checkout & Confirmation", "~1:30"),
    ("14", "Rizza", "Closing / Q&A",           "~0:30"),
    ("",   "TOTAL", "",                        "~13:45"),
]

FILL_ROW = {
    "Rizza": "EAF2ED", "Carla": "EFF6FF",
    "Julio": "FFFBEB", "Lacel": "F5F3FF", "TOTAL": "F3F4F6",
}

for row_data in TIMING:
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, row_data)):
        cell.text = val
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        presenter = row_data[1]
        fill_hex = FILL_ROW.get(presenter, "FFFFFF")
        if row_data[1] == "TOTAL":
            run.bold = True
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(nqn('w:val'), 'clear'); shd.set(nqn('w:color'), 'auto'); shd.set(nqn('w:fill'), fill_hex)
        tcPr.append(shd)

blank(10)

# ── TIPS ──────────────────────────────────────────────────────────
shaded_para("TIPS FOR THE DAY", fill_hex="2D2D2D", text_color=WHITE, size=11)

tips = [
    ("Rizza", "Keep track of the time. Signal Lacel if you're approaching 13 minutes."),
    ("Lacel", "Open boxnijuanph.vercel.app before the presentation starts. Sign in with your test account so the My Box dashboard is ready to show."),
    ("Julio", "On Slide 7, point out \"up to 12 items\" for the Custom plan — this was previously described as \"unlimited\" and has been corrected across all docs and the live app."),
    ("Carla", "If the professor asks about CSR — emphasize that the count is real-time and embedded at every step of the purchase flow, not just a landing page badge."),
    ("All",   "Q&A expected: \"Why localStorage?\" — answer: RA 10173 compliance, prototype scope. Our recommendation for future development is Supabase or Firebase, documented in our conclusions section."),
]

for name, tip_text in tips:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(10)
    r1 = p.add_run(f"{name}:  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = TEXT_MAP.get(name.upper(), DARK)
    r2 = p.add_run(tip_text)
    r2.font.size = Pt(10); r2.font.color.rgb = DARK

blank(10)

# Footer
divider("5F8F72", 4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("BoxNiJuanPH  ·  BSITOUMN COMP 047  ·  PUP Open University System  ·  June 2026")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY

out = r"C:\Users\I331629\Documents\BoxNiJuanPH-guide\BoxNiJuanPH_PresentationScript_June21.docx"
doc.save(out)
print("Done:", out)
