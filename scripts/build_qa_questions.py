"""
Generates BoxNiJuanPH_PossibleQuestions.docx
Possible professor/panel questions + suggested answers for the oral defense.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = section.bottom_margin = int(1134)
    section.left_margin = section.right_margin = int(1134)

# ── Styles ────────────────────────────────────────────────────────
GREEN = RGBColor(0x5F, 0x8F, 0x72)
DARK  = RGBColor(0x2D, 0x2D, 0x2D)
GRAY  = RGBColor(0x6B, 0x72, 0x80)
AMBER = RGBColor(0x92, 0x40, 0x0E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run("BoxNiJuanPH")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Possible Defense Questions & Suggested Answers")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BSITOUMN COMP 047  ·  PUP Open University System  ·  June 21, 2026")
    r.font.size = Pt(10); r.italic = True; r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by: Centeno, Lacel Carolane L.  (Project Lead, UI/UX Design)")
    r.font.size = Pt(10); r.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(20)

    # Divider
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '5F8F72')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # Green shaded background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAF2ED')
    pPr.append(shd)
    r = p.add_run("  " + text + "  ")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = GREEN

def add_q(number, question, answer, tip=None):
    # Question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Q{number}.  {question}")
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK

    # Answer
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("Suggested Answer:  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = GREEN
    r2 = p.add_run(answer)
    r2.font.size = Pt(10); r2.font.color.rgb = DARK

    # Tip
    if tip:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run("💡 Tip:  ")
        r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = AMBER
        r2 = p.add_run(tip)
        r2.font.size = Pt(9.5); r2.italic = True; r2.font.color.rgb = AMBER

    # Light divider
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '2')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Build Document ─────────────────────────────────────────────────
add_cover()

# ── SECTION 1: About the Project ──────────────────────────────────
add_section_heading("01 — About the Project")

add_q(1,
    "What exactly is BoxNiJuanPH? Can you describe it in one sentence?",
    "BoxNiJuanPH is a live, web-based e-commerce platform where Filipino consumers can build and order their own personalized monthly wellness box — choosing from 20 products across 4 categories.",
    "Keep it short and confident. One sentence is enough."
)

add_q(2,
    "Why wellness products? Why not a different niche?",
    "The Philippine wellness market is one of the fastest-growing consumer segments, per PSA 2023. Health-conscious Filipinos aged 18–35 are the target demographic, and there is currently no locally-focused, customizable wellness subscription platform serving this market.",
    "Mention PSA 2023 and the gap in the market."
)

add_q(3,
    "What does 'PH' in BoxNiJuanPH stand for? Why 'Juan'?",
    "'Juan' refers to Juan dela Cruz — the Filipino everyman. 'PH' anchors the brand in the Philippines. The name means 'Juan's Box in the Philippines' — a box built for every Filipino.",
)

add_q(4,
    "Is this a real business or just a prototype?",
    "It is a fully functional prototype — the platform is live, deployed, and accessible at boxnijuanph.vercel.app. All features work end-to-end. The only simulated parts are the payment processing and the OAuth authentication, which would be replaced by real integrations in a production version.",
    "Don't say 'fake' — say 'simulated' or 'prototype-level'."
)

# ── SECTION 2: Technical Questions ────────────────────────────────
add_section_heading("02 — Technical Questions")

add_q(5,
    "Why did you use Next.js instead of a simpler framework like plain HTML or React?",
    "Next.js gives us server-side rendering, file-based routing, and built-in optimizations like image optimization and code splitting — all out of the box. It also deploys seamlessly to Vercel, which is made by the same team. For a multi-page application like ours, Next.js was the most appropriate choice.",
)

add_q(6,
    "Why localStorage? Isn't that insecure or limited?",
    "For a prototype of this scope, localStorage is appropriate. All data stays on the user's device — nothing is sent to an external server — which actually makes it more privacy-compliant with RA 10173. The limitation is that data doesn't sync across devices, which we acknowledge in our Scope and Limitations. For production, we recommend Supabase or Firebase.",
    "Mention RA 10173 compliance — it frames localStorage as a privacy feature, not a weakness."
)

add_q(7,
    "Why is the authentication simulated? Why not real Google login?",
    "Real OAuth requires a registered Google Cloud project, a backend server to handle tokens, and environment secrets — all of which are beyond the scope of a student prototype and would introduce real privacy implications. Our simulated social login replicates the exact UI and UX of real OAuth without storing actual credentials.",
)

add_q(8,
    "What happens to user data if the browser is cleared?",
    "All data is lost, which is a known limitation of localStorage. We disclose this in the guest checkout warning and in our privacy policy. It is also listed as a recommendation for future enhancement — replacing localStorage with a real database.",
)

add_q(9,
    "How did you handle RA 10173 compliance?",
    "We implement RA 10173 in four ways: a full 8-section Privacy Policy page, a privacy disclosure at signup, checkout, and the contact form, no personal data is transmitted to any server, and users have full control over their data through the My Box profile tab.",
)

add_q(10,
    "Why Tailwind CSS v4? That's a newer version — what changed?",
    "Tailwind CSS v4 removes the need for a separate tailwind.config.js file and uses CSS-native variables for theming. This made our design system — the sage green palette, spacing, and typography — easier to maintain directly in CSS.",
    "If you don't remember the exact v4 differences, just say: 'v4 uses native CSS variables for theming, which made our design system easier to configure.'"
)

add_q(11,
    "What is the total number of components or pages in the project?",
    "The platform has 11 pages and multiple reusable components — including the ChatBot widget, the Step Indicator, the Variant Modal, and the Wellness Quiz. All components are written in TypeScript with strict typing.",
)

# ── SECTION 3: Features & Design ──────────────────────────────────
add_section_heading("03 — Features & Design Questions")

add_q(12,
    "Why does the Custom plan only go up to 12 items and not truly unlimited?",
    "We initially described it as 'unlimited' in early documentation, but we updated this to 'up to 12 items' to accurately reflect the platform's actual behavior. A practical upper limit of 12 items ensures the box builder performs well and sets realistic expectations for users. At ₱1,299/month, 12 items is already significantly more generous than any other tier.",
    "This is a likely question. Be direct — say you corrected it from 'unlimited' to 'up to 12 items' after refining the prototype."
)

add_q(13,
    "How does the Wellness Quiz work technically?",
    "It is a 3-step client-side quiz with state managed using React's useState hook. The user's answers to the three questions are passed into a getRecommendation function, and the third answer — how many items they want — maps directly to one of the four plan tiers: Basic, Standard, Premium, or Custom.",
)

add_q(14,
    "What does the ChatBot (BoxBot) actually do? Is it AI?",
    "BoxBot is a rule-based chatbot — it responds to 10 keyword categories using pattern matching. It is not AI in the machine learning sense. The responses are pre-written and triggered by keywords like 'refund', 'cancel', 'delivery', or 'plans'. For a prototype, this covers the most common support queries effectively.",
)

add_q(15,
    "What is the refund policy in the platform?",
    "Users can request a refund or replacement within 7 days of delivery for damaged, incorrect, or missing items. They email support@boxnijuanph.com with their order number and a photo of the issue. This is surfaced in four places: the Order Confirmation page, the My Box manage section, the Contact page topic pre-fill, and the ChatBot refund response.",
)

add_q(16,
    "Why did you include a Testimonials section? The users aren't real.",
    "The testimonials are sample data for prototype demonstration purposes — standard practice in UX prototyping. They demonstrate what the section would look like in a real deployment and provide social proof design patterns that the professor can evaluate. We acknowledge they are sample data.",
)

add_q(17,
    "How does your savings badge work on the Plans page?",
    "Each subscription plan has a defined one-time equivalent price. The savings badge computes the difference: for example, the Basic plan is ₱399/month as a subscription versus ₱549 if bought one-time, so the badge shows 'Save ₱150'. This is a common e-commerce pattern used to highlight subscription value.",
)

# ── SECTION 4: CSR & SDG ──────────────────────────────────────────
add_section_heading("04 — CSR & SDG 12 Questions")

add_q(18,
    "How exactly is UN SDG 12 integrated into the platform?",
    "SDG 12 is about Responsible Consumption and Production. Our integration is threefold: First, users choose only what they want — no unwanted items, no waste. Second, the platform highlights eco-friendly products with an Eco badge. Third, the CSR impact count at Box Summary, Order Confirmation, and My Box shows users in real time how many local and eco-friendly items they chose. It is embedded in the purchase flow, not just a marketing statement.",
)

add_q(19,
    "Are the Filipino Brand and Eco badges verified by any external body?",
    "No — for the prototype, the badges are manually assigned in our product data based on our research. In a real production platform, these would be verified through supplier documentation and third-party certifications. We acknowledge this as a limitation.",
)

add_q(20,
    "How does buying from BoxNiJuanPH actually support local brands?",
    "Every product tagged as 'Filipino Brand' in our catalog is sourced from a local Philippine wellness brand. When a user selects that product, the platform highlights it with the badge — creating awareness and driving intentional purchasing behavior. The impact summary reinforces this at the end of every order.",
)

# ── SECTION 5: Business & Market ──────────────────────────────────
add_section_heading("05 — Business & Market Questions")

add_q(21,
    "How is BoxNiJuanPH different from Good Box PH?",
    "Good Box PH is fully pre-curated — users have no say in what goes inside their box. BoxNiJuanPH gives users full item-level selection on every plan tier. We also have a subscription management dashboard, a privacy policy, a contact form, and a ChatBot — features that Good Box PH does not have.",
)

add_q(22,
    "What is the total development cost of this project?",
    "Zero pesos. All tools used are free: Next.js, TypeScript, Tailwind CSS v4, Vercel Hobby Plan, GitHub, Unsplash for product images, and Google Fonts. We had no hosting fees, no paid APIs, and no paid software.",
)

add_q(23,
    "What are your recommendations for future development?",
    "We have eight recommendations in our conclusion: integrate PayMongo or GCash API for real payments, add a Supabase or Firebase backend for persistent data, implement real OAuth, create a live product inventory system, integrate logistics APIs like LBC or J&T for real delivery tracking, expand the product catalog, add SEO optimization, and conduct formal user testing with real respondents.",
)

add_q(24,
    "Why is your target market limited to Metro Manila for delivery?",
    "Delivery simulation is scoped to Metro Manila as a reasonable starting point for a local wellness brand prototype. In a real deployment, logistics coverage would expand through courier partnerships. We acknowledge this as a scope limitation.",
)

# ── SECTION 6: Process & Academic ─────────────────────────────────
add_section_heading("06 — Process & Academic Questions")

add_q(25,
    "How did you divide the work among group members?",
    "Lacel handled project lead, UI/UX design, and front-end development. Carla managed content strategy and product catalog research. Ivy handled market research and CSR documentation. Julio was responsible for quality assurance and use case testing. Rizza handled documentation and presentation design.",
    "Mention Ivy's absence: 'Ivy is absent today due to a power interruption in her area, but she contributed to the market research and CSR sections.'"
)

add_q(26,
    "Did you conduct actual user testing?",
    "We conducted internal group testing using our QA checklist, which covers all 200+ test items across every page and feature. Formal user testing with external respondents is listed as a recommendation for future work.",
)

add_q(27,
    "Why did you choose a subscription box as your e-commerce concept?",
    "The subscription box model is one of the fastest-growing e-commerce categories globally — at 18.3% CAGR per Grand View Research 2023. It also gave us an opportunity to implement a full user journey: plan selection, box building, checkout, subscription management, and cancellation. This complexity made it a meaningful technical challenge for COMP 047.",
)

add_q(28,
    "What was the hardest part of building this?",
    "The box builder — specifically managing item selection state across plan tiers, handling the Custom plan's different flow, and ensuring that all selected items persist correctly through Summary, Checkout, and into the My Box dashboard. State management across pages without a backend required careful use of localStorage and data structure design.",
    "This is a great question to answer honestly — shows technical depth."
)

# ── FOOTER ────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BoxNiJuanPH  ·  BSITOUMN COMP 047  ·  PUP Open University System  ·  June 2026")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY

out = r"C:\Users\I331629\Documents\BoxNiJuanPH-guide\BoxNiJuanPH_PossibleQuestions.docx"
doc.save(out)
print("Done:", out)
